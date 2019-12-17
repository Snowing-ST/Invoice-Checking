# -*- coding: utf-8 -*-
"""
Created on Sat Oct 12 16:55:11 2019

解决：
1. 出现输入验证码后，“查验”为灰色，或呈蓝色但点击后“验证码失效”，说明发票查验次数太多，需刷新重输
2. 发票查验太频繁，请等1分钟
3. base64编码问题



@author: situ
"""

from selenium import webdriver
from selenium.common.exceptions import StaleElementReferenceException
from selenium.common.exceptions import ElementClickInterceptedException
#from selenium.webdriver.chrome.options import Options
import time
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
#import sys
from base64 import b64decode
import base64
from matplotlib import pyplot as plt
from cv2 import imread,split,merge

    

def yzm(driver,invoice_num,yzm_img_path):
#    img_url = driver.find_element_by_id('yzm_img').get_attribute("src")
#    urllib.request.urlretrieve(img_url,screenshot_save_path +'yzmimages.png')
#    Image.open(screenshot_save_path +'yzmimages.png') 
    print("正在查验发票%s..." % invoice_num)
    img_str = driver.find_element_by_id("yzm_img").get_attribute('src')
    img_str = img_str.split(",")[-1] # 删除前面的 “data:image/jpeg;base64,”
    img_str = img_str.replace("%0A", '\n')  # 将"%0A"替换为换行符
    

    img_data = b64decode(img_str)
    
    with open(os.path.join(yzm_img_path,'%s.jpg'%invoice_num), 'wb') as fout:
        fout.write(img_data)
        fout.close()
    #输出验证码图片
    im = imread(os.path.join(yzm_img_path,'%s.jpg'%invoice_num))     
    
    if im is None:
        print("请自行在浏览器中查看验证码")
    else:
        b,g,r=split(im)
        im2 = merge([r,g,b])
        plt.imshow(im2) 
        plt.xticks([]),plt.yticks([]) #隐藏坐标线 
        plt.show() #显示出来
    
    #输出提示文字
    text = driver.find_element_by_id("yzminfo").text
    dic = {"黄色":"33","红色":"31","蓝色":34}
    if text=="请输入验证码文字":
        print(text)
    else:
        print('请输入验证码图片中\033[0;%sm%s\033[0m文字'%(dic[text[9:11]],text[9:11]))

    code = input("请输入验证码:")
    driver.find_element_by_id('yzm').send_keys(code)
    time.sleep(1)
    driver.find_element_by_id('checkfp').click()
    time.sleep(4)
    try:
        driver.find_element_by_id('fpdm')
#        print("未跳转，等待刷新")
        raise Exception("未跳转，等待刷新")
    except:
        print("已跳转")
        pass

def screen_shot(driver,screenshot_save_path,invoice_num):
    
    time.sleep(3)
    #鼠标滚动至浏览器最上方
    driver.execute_script("""
    (function () {
        var y = 0;
        var step = 100;
        window.scroll(0, 0);
       
    })();
""")  


    driver.save_screenshot(os.path.join(screenshot_save_path,'%s.png' % invoice_num))
    print("发票%s截图成功" % invoice_num)
    time.sleep(2)
    
    
def invoice_check(driver,invoice_code,invoice_num,date,value,screenshot_save_path,yzm_img_path):
    time.sleep(2)
    driver.get("https://inv-veri.chinatax.gov.cn/")
    
    driver.find_element_by_id('fpdm').send_keys(invoice_code) 
    driver.find_element_by_id('fphm').send_keys(invoice_num)
    driver.find_element_by_id('kprq').send_keys(date)
    driver.find_element_by_id('kjje').send_keys(value)
    
    try:
        yzm(driver,invoice_num,yzm_img_path)
        false = 0
    except Exception:
        false = 1 
    except:
        print("base64图片获取有误，请等待刷新")
        false = 1
        
    while false:
        driver.refresh()
        driver.find_element_by_id('fpdm').send_keys(invoice_code) 
        driver.find_element_by_id('fphm').send_keys(invoice_num)
        driver.find_element_by_id('kprq').send_keys(date)
        driver.find_element_by_id('kjje').send_keys(value)
        try:
            yzm(driver,invoice_num,yzm_img_path)
            false = 0
        except ElementClickInterceptedException as msg:
            try:
                if driver.find_element_by_id('popup_message').text=="验证码请求次数过于频繁，请1分钟后再试！":
                    print("验证码请求次数过于频繁，1分钟后程序将自动跳转1")
                    popup = driver.find_element_by_css_selector('#popup_ok')
                    popup.click()
                    time.sleep(60)
                
            except:
                print("no1")
            false = 1
                
            
        except base64.binascii.Error as msg:
            try:
                if driver.find_element_by_id('popup_message').text=="验证码请求次数过于频繁，请1分钟后再试！":
                    print("验证码请求次数过于频繁，1分钟后程序将自动跳转2") #起作用的是这个
                    popup = driver.find_element_by_css_selector('#popup_ok')
                    popup.click()
                    time.sleep(60)
            except:
                print("no2")
                print("base64图片获取有误，请等待刷新%s"%msg)
            false = 1


    try:
        popup = driver.find_element_by_css_selector('#popup_ok')
        print("验证码输入错误，请重输!")
    except:
        popup = None
        screen_shot(driver,screenshot_save_path,invoice_num)

    while popup:  # 输错后弹窗，点击换验证码，反复输入
        popup.click()
        
        driver.find_element_by_css_selector('#yzm_img').click()#刷新验证码
#        driver.find_element_by_id('yzm_img').click()
        driver.find_element_by_id('yzm').clear()
        time.sleep(1)
        
        yzm(driver,invoice_num,yzm_img_path)
        
        try:
            popup = driver.find_element_by_css_selector('#popup_ok')
            print("验证码输入错误，请重输!")
        except:
            screen_shot(driver,screenshot_save_path,invoice_num)
           
    


def main():


#    invoice_file_path = input("请输入发票信息excel路径：")
    invoice_file_path = "E:/self_programming/invoice_check_2.0/invoice_sample_test.xlsx"
    path = os.path.dirname(invoice_file_path)

    
    os.chdir(path)
    
    screenshot_save_path = os.path.join(path,"截图")
    doc_save_path = os.path.join(path,"文档")
    yzm_img_path = os.path.join(path,"captcha")
    if not os.path.exists(screenshot_save_path):
        os.mkdir(screenshot_save_path) 
    if not os.path.exists(doc_save_path):
        os.mkdir(doc_save_path) 
    if not os.path.exists(yzm_img_path):
        os.mkdir(yzm_img_path) 
    
    invoice_info = pd.read_excel(invoice_file_path,dtype="str")#这个会直接默认读取到这个Excel的第一个表单
    invoice_info.head()
    
#    invoice_info.drop(list(range(8)),inplace=True)
#    invoice_info.reset_index(level=None, drop=True,inplace=True)
#    invoice_info.head()
    
    
    driver = webdriver.Chrome(os.path.join(path,"chromedriver.exe"))
    driver.maximize_window()
    
    for i in range(len(invoice_info)):
        try:
            invoice_code,invoice_num,date,value = [str(ele) for ele in invoice_info.ix[i,1:5]]
            invoice_check(driver,invoice_code,invoice_num,date,value,screenshot_save_path,yzm_img_path)
            time.sleep(1)
        except StaleElementReferenceException as msg:
            print("查找元素异常%s，请等待..."%msg)
            print("重新获取元素...")
            driver.refresh()
            i = i+1
        except ElementClickInterceptedException as msg:
            try:
                if driver.find_element_by_id('popup_message').text=="验证码请求次数过于频繁，请1分钟后再试！":
                    print("验证码请求次数过于频繁，1分钟后程序将自动跳转3")
                    popup = driver.find_element_by_css_selector('#popup_ok')
                    popup.click()
                    time.sleep(60)
            except:
                print("no3")
            driver.refresh()
            i = i+1                

#            invoice_code,invoice_num,date,value = [str(ele) for ele in invoice_info.ix[i,:]]
#            invoice_check(driver,invoice_code,invoice_num,date,value,screenshot_save_path)
#            time.sleep(1)
#    driver.close()

        
    print("截图已完成!请打开%s查看，正在将截图插入word文档......" % screenshot_save_path)
    
#将同一笔业务的截图放入一个word文档中，以业务编号命名
    for business_no in invoice_info["业务编号"].value_counts().index:
        invoice_no = invoice_info["发票号码"][invoice_info["业务编号"]==business_no]
        doc = Document()    # doc对象
        doc.add_heading(business_no ,0)
    
        for invoice_no_i in invoice_no:
            #string = '文字内容'
            images = os.path.join(screenshot_save_path,'%s.png' % invoice_no_i)    # 保存在本地的图片
            
            #doc.add_paragraph(string)   # 添加文字
            doc.add_picture(images,width=Inches(6))     # 添加图, 设置宽度
        doc.save(os.path.join(doc_save_path,'%s.docx' % business_no))     # 保存路径   
    
    print("已将截图插入word文档，请打开%s查看!" % doc_save_path)
    #清空验证码缓存图片
#    import shutil
#    shutil.rmtree(yzm_img_path)    #递归删除文件夹


if __name__ == "__main__":
    main()
    os.system("pause")


#pyinstaller -p C:/Users/situ/Anaconda2/envs/py3/Lib/site-packages -D check_invoice.py
#cd E:/self_programming/invoice_check_test/dist/check_invoice
    
 
#invoice_code = "5300193130"
#invoice_num = "01028425"
#date = "20191211"
#value = "877505.31"
#    
#driver = webdriver.Chrome(os.path.join(path,"chromedriver.exe"))
#driver.maximize_window()
#driver.get("https://inv-veri.chinatax.gov.cn/")
#
#driver.find_element_by_id('fpdm').send_keys(invoice_code) 
#driver.find_element_by_id('fphm').send_keys(invoice_num)
#driver.find_element_by_id('kprq').send_keys(date)
#driver.find_element_by_id('kjje').send_keys(value)
#
#invoice_check(driver,invoice_code,invoice_num,date,value,screenshot_save_path,yzm_img_path)
#driver.refresh()
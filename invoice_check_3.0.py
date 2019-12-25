# -*- coding: utf-8 -*-
"""
Created on Tue Dec 17 17:20:46 2019
gui
类的形式

@author: situ
"""

from selenium import webdriver
from selenium.common.exceptions import StaleElementReferenceException
from selenium.common.exceptions import ElementClickInterceptedException
#from selenium.webdriver.chrome.options import Options
import time
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
#import sys
from base64 import b64decode
import base64

import tkinter as tk  # 使用Tkinter前需要先导入
from tkinter import filedialog
from PIL import Image, ImageTk


class invoice_checking(tk.Tk):
    def __init__(self):
        super().__init__()
        

        
#        self.window = tk.Tk()
#        self.window.title('发票查验小助手')  
#        self.window.geometry('500x300') 
        self.title("发票查验小助手")
        self.geometry('500x300') 
        self.input_yzm = tk.StringVar()     
        self.yzm_info = tk.StringVar()
        self.i = 0 #正在查询第几个发票
        self.createWidgets()

        
#        self.iconbitmap('bitbug_favicon.ico')
#        self.window.mainloop()  
        
        
    def createWidgets(self):
        def file_input():
            filename = filedialog.askopenfilename(title='导入excel文件')
            entry_filename.insert('insert', filename) 
            self.path = os.path.dirname(entry_filename.get())
            self.screenshot_save_path = os.path.join(self.path,"截图")
            self.doc_save_path = os.path.join(self.path,"文档")
            self.yzm_img_dir = os.path.join(self.path,"captcha")
            self.invoice_info = pd.read_excel(entry_filename.get(),dtype="str")
            
        menubar = tk.Menu(self)     
        filemenu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label='File', menu=filemenu)
        filemenu.add_command(label='Open', command=file_input)
        self.config(menu=menubar)
        
        entry_filename = tk.Entry(self, width=30,font=("arial", 10, 'bold'))
        entry_filename.grid(column=0, row=0)

        b1 = tk.Button(self, text="开始",command=self.begin)
        b1.grid(column=1, row=0)
#        b1.configure(state='disabled')    # Disable the Button Widge
    
        
        self.e_yzm = tk.Entry(self, textvariable=self.input_yzm,show=None, width=10, font=('Arial', 14))  # 显示成明文形式
        self.e_yzm.grid(column=0, row=2)    
        self.e_yzm.focus()   # Place cursor(光标) into name Entry
        self.e_yzm.bind("<Return>", self.submit)  
        
        
        self.show_info = tk.Label(self, textvariable = self.yzm_info, bg='white', fg='black', font=('Microsoft YaHei', 10), width=20, height=1)
        self.show_info.grid(column=1, row=2)     
        
        self.show_yzm = tk.Label(self,fg = "black",bg = "white")
#        self.show_yzm = tk.Canvas(self,width = 90,height = 35,bg='white')
        self.show_yzm.grid(column=0, row=3)        
        
        b2 = tk.Button(self, text="查验",command=self.submit, default='active')
        b2.grid(column=0, row=4)

        
    def try_once_invoice_check(self):
        try:
            self.invoice_code,self.invoice_num,self.date,self.value = [str(ele) for ele in self.invoice_info.ix[self.i,["发票代码","发票号码","发票日期","税前金额"]]]
            self.invoice_check()
            time.sleep(1)
        except StaleElementReferenceException as msg:
            print("查找元素异常%s，请等待..."%msg)
            print("重新获取元素...")
            self.driver.refresh()
            self.i = self.i+1        
        
        
    def begin(self):
        if not os.path.exists(self.screenshot_save_path):
            os.mkdir(self.screenshot_save_path) 
        if not os.path.exists(self.doc_save_path):
            os.mkdir(self.doc_save_path) 
        if not os.path.exists(self.yzm_img_dir):
            os.mkdir(self.yzm_img_dir)
            
        self.driver = webdriver.Chrome(os.path.join(self.path,"chromedriver.exe"))
        self.driver.maximize_window()
        
        self.try_once_invoice_check()#改循环为单次

        
    
    def invoice_check(self):
        time.sleep(2)
        self.driver.get("https://inv-veri.chinatax.gov.cn/")
        try: #如果谷歌浏览器提示非私密链接
            self.driver.find_element_by_id('fpdm')
        except:
            self.driver.find_element_by_id("details-button").click()
            time.sleep(1)
            self.driver.find_element_by_id('proceed-link').click()
            time.sleep(2)        
        self.driver.find_element_by_id('fpdm').send_keys(self.invoice_code) 
        self.driver.find_element_by_id('fphm').send_keys(self.invoice_num)
        self.driver.find_element_by_id('kprq').send_keys(self.date)
        self.driver.find_element_by_id('kjje').send_keys(self.value)  
        
        time.sleep(2)
        
        try:
            self.yzm()
            false = 0
        except Exception:
            false = 1 
        except:
            print("base64图片获取有误，请等待刷新")
            false = 1
            
        while false:
            self.driver.refresh()
            self.driver.find_element_by_id('fpdm').send_keys(self.invoice_code) 
            self.driver.find_element_by_id('fphm').send_keys(self.invoice_num)
            self.driver.find_element_by_id('kprq').send_keys(self.date)
            self.driver.find_element_by_id('kjje').send_keys(self.value)
        
            time.sleep(6)
        
            try:
                self.yzm()
                false = 0
                    
                
            except base64.binascii.Error as msg:
                try:
                    if self.driver.find_element_by_id('popup_message').text=="验证码请求次数过于频繁，请1分钟后再试！":
                        print("验证码请求次数过于频繁，1分钟后程序将自动跳转2") #起作用的是这个
                        popup = self.driver.find_element_by_css_selector('#popup_ok')
                        popup.click()
                        time.sleep(60)
                except:
                    print("no2")
                    print("base64图片获取有误，请等待刷新%s"%msg)
                false = 1
    
    

    
    def yzm(self):
        print("正在查验发票%s..." % self.invoice_num)
        img_str = self.driver.find_element_by_id("yzm_img").get_attribute('src')
        img_str = img_str.split(",")[-1] # 删除前面的 “data:image/jpeg;base64,”
        img_str = img_str.replace("%0A", '\n')  # 将"%0A"替换为换行符
    
        img_data = b64decode(img_str)
        
        with open(os.path.join(self.yzm_img_dir,'%s.jpg'%self.invoice_num), 'wb') as fout:
            fout.write(img_data)
            fout.close()
        #输出验证码图片
        img_open = Image.open(os.path.join(self.yzm_img_dir,'%s.jpg'%self.invoice_num))  
        img_resize = img_open.resize((120,45),Image.ANTIALIAS)
        img = ImageTk.PhotoImage(img_resize)
        
        #label
        self.show_yzm.config(image=img)  
        self.show_yzm.image=img #keep a reference 
        # canvas
#        self.show_yzm.create_image(45,17,anchor = "n",image=img)


        #输出提示文字
        dic = {"黄色":"yellow","红色":"red","蓝色":"blue"}
        text = self.driver.find_element_by_id("yzminfo").text
        

  
        self.yzm_info.set(text)
        if text=="请输入验证码文字":
            self.show_info.config(fg = "black")
            
        else:
            self.show_info.config(fg = dic[text[9:11]])

        
        

    def submit(self,ev = None):#enter键提交
#提交验证码   
        self.driver.find_element_by_id('yzm').send_keys(self.e_yzm.get())
        time.sleep(1)
        self.driver.find_element_by_id('checkfp').click()
        self.e_yzm.delete(0,tk.END)
        time.sleep(2)
#        try:
#            self.driver.find_element_by_id('fpdm')
#            print("未跳转，等待刷新")
##            raise Exception("未跳转，等待刷新")
#        except:
#            print("已跳转")
#            pass            
       
        
        try:
            popup = self.driver.find_element_by_css_selector('#popup_ok')
            print("验证码输入错误，请重输!")
            popup.click()
            
            self.driver.find_element_by_css_selector('#yzm_img').click()#刷新验证码
    #        driver.find_element_by_id('yzm_img').click()
            self.driver.find_element_by_id('yzm').clear()
            time.sleep(1)
            
            self.yzm()

        except:
            popup = None
            time.sleep(2)
            self.screen_shot()
            self.i=self.i+1
            self.next()
    

        
    def screen_shot(self):
        
        time.sleep(1)
        #鼠标滚动至浏览器最上方
        self.driver.execute_script("""
        (function () {
            var y = 0;
            var step = 100;
            window.scroll(0, 0);
           
        })();
    """)  
    
    
        self.driver.save_screenshot(os.path.join(self.screenshot_save_path,'%s.png' % self.invoice_num))
        print("发票%s截图成功" % self.invoice_num)
        time.sleep(2)
        
    def next(self):
        if self.i<len(self.invoice_info):
            self.try_once_invoice_check()
            
        if self.i==len(self.invoice_info):
            print("截图已完成!请打开%s查看，正在将截图插入word文档......" % self.screenshot_save_path)
            
        #将同一笔业务的截图放入一个word文档中，以业务编号命名
            for business_no in self.invoice_info["业务编号"].value_counts().index:
                invoice_no = self.invoice_info["发票号码"][self.invoice_info["业务编号"]==business_no]
                doc = Document()    # doc对象
                doc.add_heading(business_no ,0)
            
                for invoice_no_i in invoice_no:
                    #string = '文字内容'
                    images = os.path.join(self.screenshot_save_path,'%s.png' % invoice_no_i)    # 保存在本地的图片
                    
                    #doc.add_paragraph(string)   # 添加文字
                    doc.add_picture(images,width=Inches(6))     # 添加图, 设置宽度
                doc.save(os.path.join(self.doc_save_path,'%s.docx' % business_no))     # 保存路径   
            
            print("已将截图插入word文档，请打开%s查看!" % self.doc_save_path)

app = invoice_checking()
app.mainloop()

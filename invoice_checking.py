# -*- coding: utf-8 -*-
"""
Created on Sat Oct 12 16:55:11 2019

@author: situ
"""

from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import StaleElementReferenceException
#from selenium.webdriver.chrome.options import Options
#import urllib
import time
import pandas as pd
#from PIL import Image
from docx import Document
from docx.shared import Inches
import os
import sys


def yzm(driver,invoice_num):
#    img_url = driver.find_element_by_id('yzm_img').get_attribute("src")
#    urllib.request.urlretrieve(img_url,screenshot_save_path +'yzmimages.png')
#    Image.open(screenshot_save_path +'yzmimages.png') 
    print("正在查验发票%s..." % invoice_num)
    code = input("请输入验证码:")
    driver.find_element_by_id('yzm').send_keys(code)
    driver.find_element_by_id('checkfp').click()
    time.sleep(3)

def screen_shot(driver,screenshot_save_path,invoice_num):
    
    time.sleep(3)
    
#    ele = driver.find_element_by_class_name('logo')
    # 将鼠标移动到定位的元素上面
#    ActionChains(driver).move_to_element(ele).perform()
#    time.sleep(2)
    driver.execute_script("""
    (function () {
        var y = 0;
        var step = 100;
        window.scroll(0, 0);

        function f() {
            if (y < document.body.scrollHeight) {
                y += step;
                window.scroll(0, y);
                setTimeout(f, 100);
            } else {
                window.scroll(0, 0);
                document.title += "scroll-done";
            }
        }

        setTimeout(f, 1000);
    })();
""")  


    driver.save_screenshot(os.path.join(screenshot_save_path,'%s.png' % invoice_num))
    print("发票%s截图成功" % invoice_num)
    time.sleep(2)
    
    
def invoice_check(driver,invoice_code,invoice_num,date,value,screenshot_save_path):
    time.sleep(2)
    driver.get("https://inv-veri.chinatax.gov.cn/")
    
    driver.find_element_by_id('fpdm').send_keys(invoice_code) 
    driver.find_element_by_id('fphm').send_keys(invoice_num)
    driver.find_element_by_id('kprq').send_keys(date)
    driver.find_element_by_id('kjje').send_keys(value)
    
    yzm(driver,invoice_num)

    try:
        popup = driver.find_element_by_css_selector('#popup_ok')
        print("验证码输入错误，请重输!")
    except:
        popup = None
        screen_shot(driver,screenshot_save_path,invoice_num)

    while popup:  # 输错后弹窗，点击换验证码，反复输入
        popup.click()
        
        driver.find_element_by_css_selector('#yzm_img').click()#刷新验证码
#        driver.find_element_by_id('yzm_img').click()
        driver.find_element_by_id('yzm').clear()
        time.sleep(1)
        
        yzm(driver,invoice_num)
        
        try:
            popup = driver.find_element_by_css_selector('#popup_ok')
            print("验证码输入错误，请重输!")
        except:
            screen_shot(driver,screenshot_save_path,invoice_num)
           
    


def main():
#    invoice_code = "5300181130"
#    invoice_num = "04414855"
#    date = "20181229"
#    value = "2927286.34"
    
    invoice_file_path = input("请输入发票信息excel路径：")
    #E:/self_programming/invoice_check_test/invoice_sample_test.xlsx
    path = os.path.dirname(invoice_file_path)

    
    os.chdir(path)
    
    screenshot_save_path = os.path.join(path,"截图")
    doc_save_path = os.path.join(path,"文档")
    if not os.path.exists(screenshot_save_path):
        os.mkdir(screenshot_save_path) 
    if not os.path.exists(doc_save_path):
        os.mkdir(doc_save_path) 
    
    invoice_info = pd.read_excel(invoice_file_path,dtype="str")#这个会直接默认读取到这个Excel的第一个表单
    invoice_info.head()
    
#    invoice_info.drop(list(range(8)),inplace=True)
#    invoice_info.reset_index(level=None, drop=True,inplace=True)
#    invoice_info.head()
    
    
    driver = webdriver.Chrome(os.path.join(path,"chromedriver.exe"))
    driver.maximize_window()
    
    for i in range(len(invoice_info)):
        try:
            invoice_code,invoice_num,date,value = [str(ele) for ele in invoice_info.ix[i,1:]]
            invoice_check(driver,invoice_code,invoice_num,date,value,screenshot_save_path)
            time.sleep(1)
        except StaleElementReferenceException as msg:
            print("查找元素异常%s，请等待..."%msg)
            print("重新获取元素...")
            driver.refresh()
            i = i+1
#            invoice_code,invoice_num,date,value = [str(ele) for ele in invoice_info.ix[i,:]]
#            invoice_check(driver,invoice_code,invoice_num,date,value,screenshot_save_path)
#            time.sleep(1)
    driver.close()
    print("截图已完成!请打开%s查看，正在将截图插入word文档......" % screenshot_save_path)
    
#将同一笔业务的截图放入一个word文档中，以业务编号命名
    for business_no in invoice_info["业务编号"].value_counts().index:
        invoice_no = invoice_info["发票号码"][invoice_info["业务编号"]==business_no]
        doc = Document()    # doc对象
        doc.add_heading(business_no ,0)
    
        for invoice_no_i in invoice_no:
            #string = '文字内容'
            images = os.path.join(screenshot_save_path,'%s.png' % invoice_no_i)    # 保存在本地的图片
            
            #doc.add_paragraph(string)   # 添加文字
            doc.add_picture(images,width=Inches(6))     # 添加图, 设置宽度
        doc.save(os.path.join(doc_save_path,'%s.docx' % business_no))     # 保存路径   
    
    print("已将截图插入word文档，请打开%s查看!" % doc_save_path)


if __name__ == "__main__":
    main()
    os.system("pause")


#pyinstaller -p C:/Users/situ/Anaconda2/envs/py3/Lib/site-packages -D check_invoice.py
#cd E:/self_programming/invoice_check_test/dist/check_invoice
    
 
